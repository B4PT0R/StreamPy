#Imports
import sys,os
_root_=os.path.dirname(os.path.abspath(__file__))
import streamlit as stl
console=stl.session_state.console
st=stl.session_state.deferrer
import traceback
import codeop
from io import BytesIO
import openai
import requests
import tiktoken
import json
from bs4 import BeautifulSoup
import PyPDF2
import docx
import odf
import subprocess
from contextlib import contextmanager
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
import shutil
 
#Tools

def root_join(*args):
    return os.path.join(_root_,*args)

from functools import reduce
def get_folder_structure(rootdir):
    '''
    Creates a nested dictionary that represents the folder structure of rootdir
    '''
    dir = {}
    rootdir = rootdir.rstrip(os.sep)
    start = rootdir.rfind(os.sep) + 1
    for path, dirs, files in os.walk(rootdir):
        folders = path[start:].split(os.sep)
        subdir = dict.fromkeys(files)
        parent = reduce(dict.get, folders[:-1], dir)
        parent[folders[-1]] = subdir
    return dir[rootdir[start:]]

def run(code):
    exec(code,globals())

def strint(i,n=3):
    s=str(i)
    while len(s)<n:
        s=str(0)+s
    return s
 
def islistoflists(l):
 
    return isinstance(l,list) and all([isinstance(e,list) for e in l])
 
def correct_code(code):
    lines=code.split('\n')
    flag=0
    for i in range(len(lines)):
        if lines[i].startswith('```') and len(lines[i])>3:
            if flag==0:
                flag+=1
                language=lines[i][3:]
                lines[i]='code=r"""'
            else:
                flag+=1
        elif lines[i]=='```':
            flag-=1
            if flag==0:
                lines[i]=f'"""\nst.code(code,language="{language}")'
        else:
            pass
    code='\n'.join(lines)
    
    lines=code.split('\n')    
    fully_correct=False
    i=0
    while not fully_correct:
        try:
            codeop.compile_command('\n'.join(lines[:i+1]),symbol='exec')
        except SyntaxError:
            lines[i]='##'+lines[i]
        else:
            i+=1
        if i==len(lines):
            fully_correct=True           
    code='\n'.join(lines)
            
    lines=code.split('\n')
    for i in range(len(lines)):
        if lines[i].startswith('##'):
            lines[i]=f'st.write("{lines[i][2:]}")'
    code='\n'.join(lines)
    
    return code
 
def isiterable(a):
 
    try:
        iter(a)
    except:
        return False
    else:
        return True
 
def get_keys(a):
 
    if isinstance(a,list):
        return range(len(a))
    elif isinstance(a,dict):
        return a.keys()
    else:
        return []
 
def treeview(startpath):
 
    s=os.path.abspath(startpath)+'/'+'\n'
    startpath = startpath.rstrip(os.sep)
    num_sep_start = startpath.count(os.sep)
    for root, dirs, files in os.walk(startpath):
        num_sep_this = root.count(os.sep)
        indent_level = num_sep_this - num_sep_start
        indent = ' ' * 4 * (indent_level)
        if not indent_level==0:
            s+='{}{}/'.format(indent, os.path.basename(root))+'\n'
        subindent = ' ' * 4 * (indent_level + 1)
        for f in files:
            s+='{}{}'.format(subindent, f)+'\n'
    return s
 
class Json:
 

    def __init__(self,json_file=None):
 
        self.data={}
        self.file=json_file
        if not os.path.exists(self.file):
            self.dump()
        else:
            self.load()        
 
            
    def load(self):
 
        with open(self.file,'r') as f:
            self.data=json.load(f)
 
    
    def dump(self):
 
        with open(self.file,'w') as f:
            json.dump(self.data,f,indent=4)
 
        
    def keys(self):
 
        return get_keys(self.data)                
 
    
    def keychains(self,keychain=[],terminal=False):
 
        item=self.read(keychain)
        key_chains = []
        if isinstance(item,dict):
            for key in get_keys(item):
                new_key_chain = keychain + [key]
                subitem=item[key]
                if isinstance(subitem, dict):
                    if not terminal:
                        key_chains.append(new_key_chain)
                    key_chains.extend(self.keychains(new_key_chain,terminal=terminal))
                else:
                    key_chains.append(new_key_chain)
        return key_chains
 
        
    def append(self,keys,value):
 
        item=self.data
        n=len(keys)
        for i in range(n):
            item=item[keys[i]]
        if isinstance(item,list):
            item.append(value)
        else:
            raise TypeError("Can only append to a list.")    
        self.dump()
 
    
    def write(self,keys,value):
 
        item=self.data
        parent=None
        n=len(keys)
        for i in range(n-1):
            if keys[i] in get_keys(item):
                parent=item
                item=item[keys[i]]
            elif isinstance(item,dict):
                item[keys[i]]=None
                parent=item
                item=item[keys[i]]
            else:
                if not item==None:
                    content=item
                    parent[keys[i-1]]={'old_data':content}
                    item=parent[keys[i-1]]
                else:
                    parent[keys[i-1]]={}
                    item=parent[keys[i-1]]
                item[keys[i]]=None
                parent=item
                item=item[keys[i]]
        if (isiterable(item) and keys[n-1] in get_keys(item)) or isinstance(item,dict):   
            item[keys[n-1]]=value
        elif isinstance(item,list) and isinstance(keys[n-1],int):
            k=keys[n-1]
            while (last_index:=(len(item)-1))<k-1:
                item.append(None)
            item.append(value)
        elif not isiterable(item):
            if not item==None:
                content=item
                parent[keys[n-2]]={'old_data':content}
                item=parent[keys[n-2]]
            else:
                parent[keys[n-2]]={}
                item=parent[keys[n-2]]
            item[keys[n-1]]=value
        else:
            raise KeyError    
        self.dump()
 
                
    def read(self,keys=[]):
 
        if len(keys)==0:
            return self.data
        else:
            item=self.data
            n=len(keys)
            for i in range(n-1):
                item=item[keys[i]]
            value=item[keys[n-1]]
            return value
 
    
    def delete(self,keys):
 
        item=self.data
        n=len(keys)
        try:
            for i in range(n-1):
                item=item[keys[i]]
            if keys[n-1] in get_keys(item):
                del item[keys[n-1]]
            self.dump()
        except:
            pass
 
     
    def clear(self):
 
        self.data={}
        self.dump()
 
    
    def __getitem__(self,key):
 
        return self.data[key]                        
 
    
    def __len__(self):
 
        return len(self.data)
 
        
    def __iter__(self):
 
        return iter(self.data)    
 
            
    def __delitem__(self,key):
 
        try:
            del self.data[key]
            self.dump()
        except KeyError:
            pass
   
def split_string(string, delimiters):
 
    """
    Splits a string into substrings according to a list of common delimiters, and then returns the list of substrings.
    The delimiters themselves should be entries of the list returned.
    """
    substrings = []
    current_substring = ""
    for char in string:
        if char in delimiters:
            if current_substring:
                substrings.append(current_substring)
                current_substring = ""
            substrings.append(char)
        else:
            current_substring += char
    if current_substring:
        substrings.append(current_substring)
    return substrings

import shutil
from selenium import webdriver
from selenium.webdriver.firefox.options import Options as FirefoxOptions
from selenium.webdriver.chrome.options import Options as ChromeOptions
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.firefox.service import Service as FirefoxService
from webdriver_manager.firefox import GeckoDriverManager
from selenium.webdriver.chrome.service import Service as ChromiumService
from webdriver_manager.core.utils import ChromeType

class browser_webdriver:

    def __init__(self):
        self.driver=None

    def __call__(self):
        if self.driver is None:
            firefox_exists = shutil.which("firefox") is not None
            chrome_exists = shutil.which("google-chrome") is not None
            chromium_exists=shutil.which("chromium") is not None
            
            # Set the options and driver based on the available browser
            if firefox_exists:
                options = FirefoxOptions()
                options.headless = True
                self.driver = webdriver.Firefox(service=FirefoxService(GeckoDriverManager().install()), options=options)
            elif chrome_exists:
                options = ChromeOptions()
                options.headless = True
                self.driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)
            elif chromium_exists:
                options = ChromeOptions()
                options.headless = True
                self.driver = webdriver.Chrome(service=ChromiumService(ChromeDriverManager(chrome_type=ChromeType.CHROMIUM).install()),options=options)
            else:
                raise Exception('No supported browser found.')
        return self.driver
    
    def close(self):
        if self.driver:
            self.driver.quit()
            self.driver = None

web_driver = browser_webdriver()

def extract_webpage_content(url):
    driver = web_driver()
    driver.get(url)
    page_source = driver.page_source
    soup = BeautifulSoup(page_source, 'html.parser')
    text = soup.get_text()
    return strip_newlines(text)
 
def get_text(source):
 
    if not isinstance(source,str):
        text=str(source)
    elif source.startswith('http'):
        url = source
        try:
            response = requests.get(url)
            content_type = response.headers['content-type']
            if content_type == 'application/pdf':
                with BytesIO(response.content) as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page_num in range(len(reader.pages)):
                        text += reader.pages[page_num].extract_text()
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                with BytesIO(response.content) as file:
                    doc = docx.Document(file)
                    text = "\n".join([para.text for para in doc.paragraphs])
            elif content_type == 'application/vnd.oasis.opendocument.text':
                with BytesIO(response.content) as file:
                    doc = odf.opendocument.load(file)
                    allparas = doc.getElementsByType(odf.text.P)
                    text='\n'.join([odf.teletype.extractText(para) for para in allparas])
            else:
                text = extract_webpage_content(url)
        except requests.exceptions.RequestException as e:
            print(f"Erreur de connexion : {e}")
            return None
    elif os.path.isfile(source):
        ext = os.path.splitext(source)[1]
        try:
            if ext == '.pdf':
                with open(source, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page_num in range(len(reader.pages)):
                        text += reader.pages[page_num].extract_text()
            elif ext == '.docx':
                doc = docx.Document(source)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == '.odt':
                doc = odf.opendocument.load(source)
                allparas = doc.getElementsByType(odf.text.P)
                text='\n'.join([odf.teletype.extractText(para) for para in allparas])
            elif ext == '.html':
                with open(source, 'r') as f:
                    soup = BeautifulSoup(f, 'html.parser')
                    text = soup.get_text()
            else:
                with open(source, 'r', encoding='utf-8') as f:
                    text = f.read()
        except FileNotFoundError:
            print(f"Fichier non trouvé : {source}")
            return None
        except Exception as e:
            print(f"Erreur lors de la lecture du fichier : {e}")
            return None
    else:
        text = source
    return strip_newlines(text)
       
def unite(list_of_lists):
    l=[]
    for liste in list_of_lists:
        l.extend(liste)
    return l
 
@contextmanager
def catch_results(console,results):
    l=[]
    n=len(console.results)
    yield
    m=len(console.results)
    l.extend(console.results[n:m])
    results.extend(unite(l))        
    
def strip_newlines(string):
    while len(string)>=1 and string[0]=='\n':
        string=string[1:]
    
    while len(string)>=1 and string[-1]=='\n':
        string=string[:len(string)-1]
    
    while not (newstring:=string.replace('\n\n', '\n'))==string:
        string=newstring
    return string
 
encoding=tiktoken.get_encoding("cl100k_base")
def token_count(string):
    return len(encoding.encode(string))
 
def split_text(text,max_tokens,overlap=50):
    tokenized_text = encoding.encode(text)
    parts = []
    current_part = []
    current_token_count = 0

    for token in tokenized_text:
        if current_token_count + 1 > max_tokens:
            parts.append(encoding.decode(current_part))
            if overlap>0:
                current_part = current_part[-overlap:]
            else:
                current_part=[]
            current_token_count = len(current_part)
        current_part.append(token)
        current_token_count += 1

    parts.append(encoding.decode(current_part))
    
    return parts
 
def pack(messages):
    text=''
    for message in messages:
        text+=message['name']+':\n'
        text+=message['content']+'\n\n'
    return text       
 
def total_tokens(messages):
    return token_count(pack(messages))

def set_OpenAI_API_key():
    openai.api_key=stl.session_state.APIkey

set_OpenAI_API_key()

def TexToPDF(tex_file, pdf_file):
 
    if tex_file.endswith('.tex') and pdf_file.endswith('.pdf'):
        if os.path.exists(tex_file):
            file_path,file_name=os.path.split(tex_file)
            try:
                subprocess.run(['pdflatex', '-interaction=nonstopmode', '-output-directory', '.', tex_file])
                temp_pdf_file = file_name.replace('.tex', '.pdf')
                os.remove(file_name.replace('.tex', '.aux'))
                os.remove(file_name.replace('.tex', '.log'))
                os.rename(temp_pdf_file, pdf_file)
            except Exception as e:
                print(str(e))
        else:
            raise Exception("tex file not found.") 
    else: 
        raise Exception("Arguments must be .tex and .pdf files paths")

class Agent:
 
    def __init__(self,console=None,utils=None):
        self.model="gpt-4"
        self.completion_mode='chat'
        self.user="Baptiste"
        self.name="Pandora"
        self.avatar=root_join('Pandora','pandora.jpeg')
        self.ID=0
        self.utils=utils
        self.type='Pandora'
        self.description="Python IDE AI assistant"
        self.OS="Lubuntu"
        self.console=console
        self.deferrer=console.deferrer
        self.mode=None
        self.modes=Json(root_join('Pandora','modes.json'))
        self.available_modules=Json(root_join('Pandora','av_modules.json'))
        self.user_prompt=''
        self.response_container=None
        self.preprompt=None
        self.reminder=None
        self.tools=None
        self.functions=None
        self.memory=None
        self.fetched_keychains=[]
        self.static=[]
        self.targets=[]
        self.triggers_target_new_turn=False
        self.auto_call=False
        self.agents={}
        self.injected=[]
        self.new_turn=False
        self.took_new_turn=False
        self.base_example=None
        self.examples=None
        self.error_count=0
        self.max_errors=2
        self.max_past=100
        self.format_preprompt=True
        self.uses_functions=False
        self.uses_python=True
        self.uses_tools=True
        self.uses_memory=True
        self.uses_static=False
        self.uses_example=True
        self.uses_past=True
        self.uses_examples=True
        self.uses_reminder=True
        self.coef=1
        self.log_mode=True
        self.learn_mode=False
        self.speech_mode=True
        self.lang='fr'
        self.max_tokens=2000
        self.token_limit=8000
        self.temperature=0.
        self.top_p=0.5
        self.n=1
        self.stream=False
        self.stop=['system:']
        self.presence_penalty=0
        self.frequency_penalty=0
        self.logit_bias=None
        self.past=[]
        self.settings=None

    def new_message(self,content,role,name):
 
        if content=="":
            return []
        else:
            message=[{"role":role,"name":name,"content":content}]
            return message    
    
    def add_to_past(self,message):
 
        self.past=self.past+message
 
    def gen_tools(self):
 
        tools='#AVAILABLE TOOLS:\n'+'\n'.join([tool for tool in self.tools])+'\n'
        tools=self.new_message(content=tools,role="system",name="system_bot")
        return tools        
    
    def gen_functions(self):
 
        functions=[self.functions[function]['for_agent'] for function in self.functions]
        return functions
    
    def add_static_item(self,content,title=None):
 
        n=len(self.static)
        if title==None:
            title='Static context item ['+str(n)+']'
        item=title+':\n'
        item+='###\n'
        item+=content+'\n'
        item+='###\n\n'
        self.static.append(item)
 
    def get_static_item(self,index):
 
        if index in get_keys(self.static):
            content=self.static[index]
        else:
            content=None
        return content
        
    def remove_static_item(self,index):
 
        if index in get_keys(self.static):
            del self.static[index]
    
    def clear_static(self):
 
        self.static=[]
        
    def gen_static(self):
 
        static=''.join(self.static)
        return self.new_message(content=static,role='system',name='system_bot')
    
    def inject(self,message):
 
        self.injected+=message
    
    def gen_injected(self):
 
        contents=self.injected.copy()
        self.injected=[]
        return contents
    
    def gen_agents(self):
 
        if not len(self.agents)==0:
            agents=['Name: '+self.agents[agent].name+' ; Type: '+self.agents[agent].type+' ; Description: '+self.agents[agent].description for agent in self.agents]
            s="#Active agents:\n"
            s+='{\n'
            s+='\n'.join(agents)
            s+='\n'
            s+='}\n'
            message=self.new_message(content=s,role='system',name='system_bot')
        else:
            message=[]
        return message
        
    def gen_preprompt(self):
 
        if self.format_preprompt:
            preprompt=self.new_message(content=self.preprompt.format(**locals()),role="system",name='system_bot')
        else:
            preprompt=self.new_message(content=self.preprompt,role="system",name='system_bot')
        return preprompt
        
    def gen_example(self):
 
        example=self.new_message(content=self.base_example,role='system',name='system_bot')   
        return example
               
    def gen_memory(self):
 
        keychains=self.fetched_keychains
        mem=''
        #mem+=f"Current settings : {str([key+'='+repr(self.settings[key]) for key in self.settings.keys()])}\n" 
        mem+=f"Available modules: {str(self.available_modules.read(['available_modules']))}\n"
        mem+=f"Fetchable memory keychains: {str(self.memory.keychains(terminal=False))}\n"
        return self.new_message(content=mem,role='system',name='system_bot')

    def gen_reminder(self):
        if not self.took_new_turn:
            return self.new_message(content="#REMINDER:\n"+self.reminder,role='system',name='system_bot')
        else:
            return [] 
                
    def gen_context(self,prompt=""):
        prompt=self.new_message(content=prompt,role='user',name=self.user) if not prompt=="" else []
        preprompt=self.gen_preprompt()
        tools=self.gen_tools() if self.uses_tools else []
        agents=self.gen_agents()
        static=self.gen_static() if self.uses_static else []
        injected=self.gen_injected()
        example=self.gen_example() if self.uses_example else []
        memory=self.gen_memory() if self.uses_memory else []
        reminder=self.gen_reminder() if self.uses_reminder else []
        context_limit=self.token_limit-self.max_tokens
        count=total_tokens(preprompt+tools+example+static+injected+memory+agents+reminder+prompt)
        n=len(self.past)
        m=len(self.examples)
        k=0
        j=0
        tokens=0
        while self.uses_past and n>=1 and k<min(n,self.max_past) and tokens+(new_tokens:=total_tokens([self.past[n-k-1]]))<=int(self.coef*(context_limit-count)):
            tokens=tokens+new_tokens
            k+=1
        past=self.past[n-k:n] if self.uses_past else []
        while self.uses_examples and m>=1 and j<m and tokens+(new_tokens:=total_tokens(self.examples[j]))<=context_limit-count:
            tokens=tokens+new_tokens
            j+=1
        examples=self.examples[:j] if self.uses_examples else []
        context=preprompt+tools+example+unite(examples)+past+static+injected+memory+agents+reminder+prompt
        context_length=tokens+count
        return context,context_length 
    
    def gen_completion_response(self,context):
 
        answer = openai.Completion.create(
            model=self.model,
            prompt=pack(context),
            max_tokens=self.max_tokens,
            temperature=self.temperature,
            top_p=self.top_p,
            n=self.n,
            stop=self.stop,
            presence_penalty=self.presence_penalty,
            frequency_penalty=self.frequency_penalty,
            #logit_bias=self.logit_bias,            
        )
        response=answer["choices"][0]
        return response

    def gen_chat_response(self,context):
 
        kwargs={
            'model':self.model,
            'messages':context,
            'max_tokens':self.max_tokens,
            'temperature':self.temperature,
            'top_p':self.top_p,
            'n':self.n,
            'stop':self.stop,
            'presence_penalty':self.presence_penalty,
            'frequency_penalty':self.frequency_penalty
        }
        
        if self.uses_functions:
            functions=self.gen_functions()
            if not len(functions)==0:
                kwargs['functions']=self.gen_functions()
                kwargs['function_call']='auto'
            
        answer = openai.ChatCompletion.create(**kwargs)
        response=answer["choices"][0]
        return response
 
    def gen_response(self,prompt):
        context,context_length=self.gen_context(prompt)      
        #edit(text=pack(context))
        success=False
        while success==False:
            try:
                if self.completion_mode=='completion':
                    response=self.gen_completion_response(context)
                elif self.completion_mode=='chat':
                    response=self.gen_chat_response(context)
            except Exception as e:
                print(str(e))
                #print(self.name+": Server error. Retrying...")
                pass
            else:
                success=True
        if response.get('message'):
            if response['message'].get('content'):
                content=response['message']['content'].strip()
            else:
                content=""
            if response['message'].get('function_call'):
                function_call=response['message']['function_call']
            else:
                function_call={}
        elif response.get('text'):
            content=response['text'].strip()
            function_call={}
        else:
            content=""
            function_call={}
    
        response={'role':'assistant','content':content,'function_call':function_call}
        return response    
    
    def process_message(self,response):
 
        if not len(response['content'])==0:
            if self.uses_python:
                response['content']=correct_code(response['content'])
            
            with self.raw_code_expander:
                self.deferrer.code(response['content'])
                
            message=self.new_message(content=response['content'],role='assistant',name=self.name)
            if self.uses_past:
                self.add_to_past(message)
            
            if self.uses_python:
                self.run(response['content'])
    
    def process_function(self,response):
        function_call=response['function_call']
        if not len(function_call)==0:
            try:
                name=function_call['name']
                kwargs=json.loads(function_call['arguments'])
                if self.functions[name]['for_system'].get('code'):
                    code=self.functions[name]['for_system']['code']
                    exec(code,globals())
                results=eval(self.functions[name]['for_system']['name'])(**kwargs)
            except Exception as e:
                results="Function call resulted in the following exception:"+str(e)
            calling=self.new_message(content="'function_call':"+json.dumps(function_call),role='assistant',name=self.name)
            results=self.new_message(content="'results':"+json.dumps(results),role='function',name=name)
            self.add_to_past(calling+results)
            self.process()
                                        
    def process(self,prompt=""):
        #self.clear_memory()
        self.new_response_container()    
        self.new_turn=False
        with self.show_status("Generating response..."):
            response=self.gen_response(prompt)                
        
        if not prompt=="":
            prompt=self.new_message(content=prompt,role='user',name=self.user)
            if self.uses_past:
                self.add_to_past(prompt)
        
        with self.show_status("Running code..."):
            self.process_message(response)
            
            self.process_function(response)
                            
        self.memory.dump()
            
        if self.new_turn:
            self.took_new_turn=True
            self.process()    
        
    def update(self):
        pass
    
    def new_response_container(self):
        #self.deferrer.mode='streamed'
        self.response_container=self.deferrer.chat_message(name='assistant',avatar=self.avatar)
        with self.response_container:
            self.status_placeholder=self.deferrer.empty()
            self.raw_code_expander=self.deferrer.expander("Raw code output.")
        
    @contextmanager
    def show_status(self,status):
        with self.status_placeholder:
            with self.deferrer.spinner(status):
                yield
     
    def __call__(self,prompt=""):
        self.took_new_turn=False
        self.user_prompt=prompt
        self.update()
        if len(self.targets)==0:
            for name in self.agents:
                agent=self.agents[name]
                if agent.auto_call:
                    agent(prompt)
            self.process(prompt)
        else:
            response=self.get_response(prompt)
            message=self.new_message(content=response,role='assistant',name=self.name)
            for target in self.targets:
                if not target.name=='Pandora':
                    target.inject(message)
                else:
                    target.add_to_past(message)
                if self.triggers_target_new_turn:
                    target.new_turn=True                                           
    
    def get_response(self,prompt=''):
 
        self.update()
        self.user_prompt=prompt
        for name in self.agents:
            agent=self.agents[name]
            if agent.auto_call:
                agent(prompt)
        response=self.gen_response(prompt)
        return response['content']
    
    def prompt(self,prompt=""):
        if prompt=="":
            pass
        else:
            with self.deferrer.chat_message(name='user'):
                self.deferrer.write(prompt)
            self(prompt)
    
    def run(self,code):
        try:
            results=[]
            with catch_results(self.console,results):
                with self.response_container:
                    exec(code,self.console.names)
                    #self.console.run(code)
            result='\n'.join(results)
        
        except Exception as e:
            self.error_count+=1
            err = traceback.format_exc()
            message="Interpreter's feedback:\n"+err
            message=self.new_message(content=message,role='system',name='system_bot')
            self.add_to_past(message)
            if self.error_count<=self.max_errors:
                self.new_turn=True
            else:
                pass
        else:
            self.error_count=0
            message="Interpreter's feedback:\n"+result
            message=self.new_message(content=message,role='system',name='system_bot')
            self.add_to_past(message)
                    
    def clear_memory(self):
        self.past=[]
        self.fetched_keychains=[]
 
    def clear_memory_file(self):
        self.memory.clear()
    
    def load_mode(self,mode):
        self.modes.load()
        self.mode=mode
        self.model=self.modes.read([mode,'model'])
        self.token_limit=self.modes.read([mode,'token_limit'])
        self.completion_mode=self.modes.read([mode,'completion_mode'])
        self.top_p=self.modes.read([mode,'top_p'])
        self.temperature=self.modes.read([mode,'temperature'])
        self.max_tokens=self.modes.read([mode,'max_tokens'])
        self.preprompt=self.modes.read([mode,'preprompt'])
        self.format_preprompt=self.modes.read([mode,'format_preprompt'])
        self.examples=self.modes.read([mode,'examples'])
        self.base_example=self.modes.read([mode,'base_example'])
        self.reminder=self.modes.read([mode,'reminder'])
        self.uses_functions=self.modes.read([mode,'uses_functions'])
        self.uses_python=self.modes.read([mode,'uses_python'])
        self.uses_past=self.modes.read([mode,'uses_past'])
        self.uses_static=self.modes.read([mode,'uses_static'])
        self.uses_examples=self.modes.read([mode,'uses_examples'])
        self.uses_example=self.modes.read([mode,'uses_example'])
        self.uses_memory=self.modes.read([mode,'uses_memory'])
        self.uses_reminder=self.modes.read([mode,'uses_reminder'])
        self.uses_tools=self.modes.read([mode,'uses_tools'])
        self.tools=self.modes.read([mode,'tools'])
        self.functions=self.modes.read([mode,'functions'])
        self.lang=self.modes.read([mode,'lang'])
        self.name=self.modes.read([mode,'name'])
        self.description=self.modes.read([mode,'description'])
        self.type=self.modes.read([mode,'type'])
        self.triggers_target_new_turn=self.modes.read([mode,'triggers_target_new_turn'])
        self.auto_call=self.modes.read([mode,'auto_call'])
        self.log_mode=self.modes.read([mode,'log_mode'])
        self.speech_mode=self.modes.read([mode,'speech_mode'])
        
    def save_mode(self,mode=None):
        self.mode=mode or self.mode
        self.modes.write([self.mode,'model'],self.model)
        self.modes.write([self.mode,'top_p'],self.top_p)
        self.modes.write([self.mode,'token_limit'],self.token_limit)
        self.modes.write([self.mode,'completion_mode'],self.completion_mode)
        self.modes.write([self.mode,'temperature'],self.temperature)
        self.modes.write([self.mode,'max_tokens'],self.max_tokens)
        self.modes.write([self.mode,'base_example'],self.base_example)
        self.modes.write([self.mode,'reminder'],self.reminder)
        self.modes.write([self.mode,'preprompt'],self.preprompt)
        self.modes.write([self.mode,'format_preprompt'],self.format_preprompt)
        self.modes.write([self.mode,'examples'],self.examples)
        self.modes.write([self.mode,'tools'],self.tools)
        self.modes.write([self.mode,'functions'],self.functions)
        self.modes.write([self.mode,'uses_functions'],self.uses_functions)
        self.modes.write([self.mode,'uses_python'],self.uses_python)
        self.modes.write([self.mode,'uses_static'],self.uses_reminder)
        self.modes.write([self.mode,'uses_reminder'],self.uses_reminder)
        self.modes.write([self.mode,'uses_tools'],self.uses_tools)
        self.modes.write([self.mode,'uses_example'],self.uses_example)
        self.modes.write([self.mode,'uses_past'],self.uses_past)
        self.modes.write([self.mode,'uses_memory'],self.uses_memory)
        self.modes.write([self.mode,'uses_examples'],self.uses_examples)
        self.modes.write([self.mode,'name'],self.name)
        self.modes.write([self.mode,'description'],self.description)
        self.modes.write([self.mode,'type'],self.type)
        self.modes.write([self.mode,'triggers_target_new_turn'],self.triggers_target_new_turn)
        self.modes.write([self.mode,'auto_call'],self.auto_call)
        self.modes.write([self.mode,'log_mode'],self.log_mode)
        self.modes.write([self.mode,'speech_mode'],self.speech_mode)
        self.modes.write([self.mode,'lang'],self.lang)
  
    def add_function(self,code):
        agent=Pandora()
        agent.load_mode('Function_Creator')
        prompt="code=\"\"\"\n"+code+"\n\"\"\""
        descriptor=json.loads(agent.get_response(prompt))
        name=descriptor['name']
        function={
            "for_agent":descriptor,
            "for_system":{
                "name":name,
                "code":code
            }
        }
        self.modes.write([self.mode,"functions",name],function)
 
    def remove_function(self,name):
        if self.modes[self.mode]["functions"].get(name):
            self.modes.delete([self.mode,"functions",name])        
 
    def gen_agent_ID(self):
        IDs=[self.agents[agent].ID for agent in self.agents]
        i=0
        while i in IDs:
            i+=1
        return i
                      
    def reduce(self,text):
        parts=split_text(text,max_tokens=2000)
        agent=Pandora()
        agent.load_mode('Summarize')
        agent.max_tokens=1000
        responses=[]
        i=1
        for part in parts:
            i+=1
            prompt=f"text=\"{part}\""
            responses.append(agent.get_response(prompt))
        response=' '.join(responses)
        return response  
 
    def get_text(self,source):
        if source in self.memory.keychains():
            text=self.assemble(source)
        else:
            text=get_text(source)
        #self.save(['_text_'],text)
        return text
              
class Pandora(Agent):
 
    def __init__(self,folder,console,utils):
        Agent.__init__(self,folder,console,utils)
        self.load_mode('Pandora-GPT4-web')
        self.memory=MemoryAgent(self.console,self.utils,self,os.path.join(folder,'Pandora','memory.json'))
        self.workfolder=folder
        if not self.workfolder in sys.path:
            sys.path.append(self.workfolder)
 
    def update(self):
        pass
             
    def get(self,keychain=[]):
        if isinstance(keychain,str):
            keychain=[keychain]
        try:
            data=self.memory.read(keychain)
            return data
        except:
            print(f"system: {str(keychain)} keychain is not referenced in memory.")
            return ''
    
    def fetch(self,keychains=[]):
        if islistoflists(keychains):
            pass
        elif isinstance(keychains,list):
            keychains=[keychains]
        elif isinstance(keychains,str):
            keychains=[[keychains]]
        fetched='Fetched data:\n'
        count=token_count(fetched)
        for keychain in keychains:
            content=str(keychain)+':\n'
            try:
                content+=str(self.get(keychain))+'\n\n'
                count+=token_count(content)
                if count>2000:
                    print("system_bot: Too much tokens in the fetched data, please use an agent to process it.")
                    break
                else:
                    fetched+=content
            except Exception as e:
                print(str(e))
        message=self.new_message(content=fetched,role='system',name='system_bot')
        self.add_to_past(message)        
        self.new_turn=True
           
    def save(self,keychain,content):
        if isinstance(keychain,str):
            keychain=[keychain]
        if not keychain[0] in ['Properties_Info','Methods_Info']:
            self.memory.write(keychain,content)
        else:
            print('system: Saving forbidden. This sector of memory is read only.')
            
    def forget(self,keychain):
        if isinstance(keychain,str):
            keychain=[keychain]
        if not keychain[0] in ['Properties_Info','Methods_Info']:
            if keychain in self.fetched_keychains:
                self.fetched_keychains.remove(keychain)
            self.memory.delete(keychain)
        else:
            print('system: Forgeting forbidden. This sector of memory is read only.') 
 
    def edit(self,file='buffer',text=None,wait=False):
        self.utils.edit(file=file,text=text,wait=wait)              
 
    def fetch_workfolder_contents(self):
        s='#Workfolder Contents:\n'
        s+=treeview(self.workfolder)
        message=self.new_message(content=s,role='system',name='system_bot')
        self.add_to_past(message)
        self.new_turn=True
 
    def text_agent(self,source,description="AI Text_Agent responding to queries about a text content passed to it."):
        text=get_text(source)
        if token_count(text)<200000:
            agent=TextAgent(self.console,self.utils,self,source,description)
            self.agents[agent.name]=agent
            return agent.name
        else:
            raise Exception('The text source is too long for a Text_Agent')
 
    def code_agent(self,source,description="AI Code_Agent responding to queries about a code content passed to it."):
        text=get_text(source)
        if token_count(text)<200000:
            agent=CodeAgent(self.console,self.utils,self,source,description)
            self.agents[agent.name]=agent
            return agent.name
        else:
            raise Exception('The code source is too long for a Code_Agent')
  
    def get_docstrings(self,python_code_or_file):
        if python_code_or_file.endswith('.py'):
            with open(python_code_or_file,'r') as f:
                code=f.read()
        else:
            code=python_code_or_file
        parts=split_text(code,max_tokens=1500)
        agent=Pandora()
        agent.load_mode('Doc_Python')
        agent.max_tokens=750
        responses=[]
        i=1
        for part in parts:
            i+=1
            prompt=f"code=\"{part}\""
            responses.append(agent.get_response(prompt))
        response='\n'.join([response for response in responses])
        lines=response.split('\n')
        i=0
        for line in lines:
            k=line.find('self')
            if not k==-1 and not line.startswith(' '):
                lines[i]='    '+lines[i]
            i+=1
        response='\n'.join(lines)
        self.save(['_docstrings_'],response)
        self.fetch([['_docstrings_']])
        return response       
 
    def assemble(self,sources,overlap=0):
 
        if sources in self.memory.keychains():
            if not sources[-1]=='table_of_content':
                contents=self.memory.read(sources)
                gathered=self.assemble(contents)
            else:
                gathered=''
        elif [sources] in self.memory.keychains():
            sources=[sources]
            if not sources[-1]=='table_of_content':
                contents=self.memory.read(sources)
                gathered=self.assemble(contents)
            else:
                gathered=''
        elif isinstance(sources,list) or isinstance(sources,dict):
            gathered=''
            for key in get_keys(sources):
                gathered+=self.assemble(sources[key])
        else:
            gathered=get_text(sources)
        return gathered        
                
    def summarize(self,sources,max_tokens=500):
        text=self.assemble(sources)
        i=1
        while token_count(text)>max_tokens:
            i+=1
            text=self.reduce(text)
        self.save(['_summarized_'],text)
        self.fetch([['_summarized_']])
        return text            
    
    def websearch(self,query,num=5,type='web'):
        results=self.utils.google_search(query,num,type)
        self.save(['_websearch_results_'],results)
        self.fetch([['_websearch_results_']])
        return results
    
    def find(self,query,sources):
        text=self.assemble(sources) 
        parts=split_text(text,max_tokens=2500)
        agent=Pandora()
        agent.load_mode('Find')
        agent.max_tokens=500
        responses=[]
        i=1
        for part in parts:
            i+=1
            prompt=f"query=\"{query}\"\ntext=\"{part}\""
            responses.append(agent.get_response(prompt))
        response='\n'.join([response for response in responses if not response=='None'])
        self.save(['_found_'],response)
        self.fetch([['_found_']])
        return response

    def tex_to_pdf(self,tex_file_or_code,pdf_file):
        TexToPDF(tex_file_or_code,pdf_file)

class MemoryAgent(Agent):
 
    def __init__(self,console,utils,parent,json_file):
        Agent.__init__(self,console,utils)
        self.load_mode('Memory_Agent')
        self.ID=parent.gen_agent_ID()
        self.format_preprompt=False
        self.auto_call=False
        self.triggers_target_new_turn=True
        self.description='An agent specialized in managing a local long-term memory json file.'
        self.targets.append(parent)
        parent.agents['Memory_Agent']=self
        self.memory=Json(json_file)

    def update(self):
        self.memory.load()
        self.clear_static()
        self.add_static_item(json.dumps(self.memory.data),title="Memory JSON file contents")
           
    def write(self,keychain,content):
        self.memory.write(keychain,content)   
 
    def keychains(self,keychain=[],terminal=False):
        return self.memory.keychains(keychain=keychain,terminal=terminal)
 
    def read(self,keychain):
        return self.memory.read(keychain)
 
    def delete(self,keychain):
        if keychain in self.keychains():
            self.memory.delete(keychain)
            
    def append(self,keychain,value):
        if keychain in self.keychains():
            self.memory.append(keychain,value)
            
    def clear(self):
        self.memory.clear()
        
    @property
    def data(self):
        return self.memory.data
    
    def load(self):
        self.memory.load()
        
    def dump(self):
        self.memory.dump()
        
    def keys(self):
        return self.memory.keys()
    
    def __getitem__(self,key):
        return self.memory[key]
        
    def __len__(self):
        return len(self.memory)
        
    def __iter__(self):
        return iter(self.memory)
        
    def __delitem__(self,key):
        if key in get_keys(self.memory.data):
            del self.memory[key]    
 
class TextAgent(Agent):
 
    def __init__(self,console,utils,parent,source,description='An AI agent specialized in dealing with long pieces of text'):
 
        Agent.__init__(self,console,utils)
        self.load_mode('Text_Agent')
        self.n=3
        self.max_tokens=int(15000/(self.n+1))
        self.ID=parent.gen_agent_ID()
        self.name='Agent_'+str(self.ID)
        self.type='Text_Agent'
        self.description=description
        text=get_text(source)
        tokens=token_count(text)
        self.targets.append(parent)
        parent.agents[self.name]=self
        if parent.type=='Pandora':
            self.triggers_target_new_turn=True
            self.auto_call=False
        else:
            self.triggers_target_new_turn=False
            self.auto_call=True
        if tokens<(15000-self.max_tokens):
            self.add_static_item(text)
        else:
            max_tokens=int(tokens/self.n)
            chunks=split_text(text,max_tokens=max_tokens,overlap=200)
            agents=[]
            for chunk in chunks:
                agent=TextAgent(console,utils,self,chunk)
                self.agents[agent.name]=agent
                    
class CodeAgent(Agent):
 
    def __init__(self,console,utils,parent,source,description='An AI agent specialized in dealing with long pieces of code'):
 
        Agent.__init__(self,console,utils)
        self.load_mode('Code_Agent')
        self.n=3
        self.max_tokens=int(15000/(self.n+1))
        self.ID=parent.gen_agent_ID()
        self.name='Agent_'+str(self.ID)
        self.type='Code_Agent'
        self.description=description
        text=get_text(source)
        tokens=token_count(text)
        self.targets.append(parent)
        parent.agents[self.name]=self
        if parent.type=='Pandora':
            self.triggers_target_new_turn=True
            self.auto_call=False
        else:
            self.triggers_target_new_turn=False
            self.auto_call=True
        if tokens<(15000-self.max_tokens):
            self.add_static_item(text)
        else:
            max_tokens=int(tokens/self.n)
            chunks=split_text(text,max_tokens=max_tokens,overlap=200)
            agents=[]
            for chunk in chunks:
                agent=CodeAgent(console,utils,self,chunk)
                self.agents[agent.name]=agent
                    
class FolderAgent(Agent):
 
    def __init__(self,console,utils,parent,folder):
 
        Agent.__init__(self,console,utils)
        self.load_mode('Folder_Agent')
        self.parent=parent
        self.workfolder=folder
        self.ID=parent.gen_agent_ID()
        self.format_preprompt=False
        self.auto_call=False
        self.triggers_target_new_turn=True
        self.description='An agent specialized in describing the contents of a local folder'
        self.targets.append(parent)
        self.parent.agents['Folder_Agent']=self
        self.update()
           
    def update(self):
 
        self.clear_static()
        content="Folder's path: "+self.workfolder+'\n'
        content+="Contents:\n"
        content+=json.dumps(get_folder_structure(self.workfolder))
        self.add_static_item(content,title="Pandora's folder contents")
        
    def get_contents(self):
 
        self.update()
        content="Folder's path: "+self.workfolder+'\n'
        content+="Contents:\n"
        content+=json.dumps(get_folder_structure(self.workfolder))
        message=self.new_message(content=content,role='assistant',name=self.name)
        self.parent.inject(message)
        self.parent.new_turn=True


